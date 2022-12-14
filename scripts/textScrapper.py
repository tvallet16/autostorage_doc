#!/usr/bin/env python
# coding: utf-8

# In[ ]:





# In[3]:


import PyPDF2
from pdfrw import PdfReader
import docx
import pandas as pd
import re
from io import StringIO


from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser



class TextScrapper():
    def __init__(self):
        pass
    
    def cleanup_remove_spaces(self, text):
        regex = re.compile(r"\s+")
        text = re.sub(regex, " ", text)
        regex = re.compile(r"\n")
        text = re.sub(regex, "", text)
        regex = re.compile(r"\t")
        text = re.sub(regex, "", text)
        return text
    
    def getTextPdf(self, filename):
        """
        text = extract_text(filename)
        return text
        """
        
        """
        output_string = StringIO()
        with open(filename, 'rb') as in_file:
            parser = PDFParser(in_file)
            doc = PDFDocument(parser)
            rsrcmgr = PDFResourceManager()
            device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            for page in PDFPage.create_pages(doc):
                interpreter.process_page(page)

        return output_string.getvalue()
        """
        
        
        
        
        fullText = []
        
        # creating a pdf file object 
        pdfFileObj = open(filename, 'rb') 
        
        # creating a pdf reader object 
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj) 

        # printing number of pages in pdf file 
        for i in range(pdfReader.numPages):
            # creating a page object 
            pageObj = pdfReader.getPage(i)

            # extracting text from page 
            pageText = pageObj.extractText()
            fullText.append(pageText)

        # closing the pdf file object 
        pdfFileObj.close() 
        return fullText
        
    def getTextCryptedPdf(self, filename):
        output_string = StringIO()
        with open(filename, 'rb') as in_file:
            parser = PDFParser(in_file)
            doc = PDFDocument(parser)
            rsrcmgr = PDFResourceManager()
            device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            for page in PDFPage.create_pages(doc):
                interpreter.process_page(page)

        return output_string.getvalue()
    
    def getTextWord(self, filename):
        doc = docx.Document(filename)
        fullText = []
        for para in doc.paragraphs:
            if para.text != '':
                fullTextClean = self.cleanup_remove_spaces(para.text)
                fullText.append(fullTextClean)
        finalText = ''.join(str(text) for text in fullText)
        return finalText

    def getTextExcel(self, filename):
        df = pd.read_excel(filename)
        print(df)
        clean_df = df.dropna()
        print(clean_df)
        return clean_df.to_string()
    
    def getTextCsv(self, filename):
        df = pd.read_csv(filename)
        return df



